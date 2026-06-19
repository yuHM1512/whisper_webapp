"""
Vietnamese Speech-to-Text Web App
Flask + faster-whisper + pyannote.audio (speaker diarization) + SSE streaming
"""

from flask import Flask, render_template, request, jsonify, Response, send_file
import os, json, threading, uuid, io
from datetime import datetime
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import queue

# Optional diarization
try:
    from pyannote.audio import Pipeline
    import torch
    DIARIZATION_AVAILABLE = True
    print("[*] pyannote.audio loaded — diarization available")
except ImportError:
    DIARIZATION_AVAILABLE = False
    print("[!] pyannote.audio not installed — diarization disabled")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

_model_cache = {}
_model_lock = threading.Lock()
_dia_pipeline = None
_dia_lock = threading.Lock()
_task_queues: dict[str, queue.Queue] = {}


def _detect_device():
    try:
        import ctranslate2
        if ctranslate2.get_cuda_device_count() > 0:
            return "cuda", "float16"
    except Exception:
        pass
    return "cpu", "int8"

_DEVICE, _COMPUTE_TYPE = _detect_device()
print(f"[*] Runtime device: {_DEVICE} / {_COMPUTE_TYPE}")


def get_model(size: str) -> WhisperModel:
    with _model_lock:
        if size not in _model_cache:
            print(f"[*] Loading Whisper model: {size} on {_DEVICE}...")
            try:
                _model_cache[size] = WhisperModel(size, device=_DEVICE, compute_type=_COMPUTE_TYPE)
            except Exception as e:
                print(f"[!] GPU load failed ({e}), falling back to CPU")
                _model_cache[size] = WhisperModel(size, device="cpu", compute_type="int8")
            print(f"[*] Whisper model {size} ready.")
        return _model_cache[size]


def get_dia_pipeline(hf_token: str):
    global _dia_pipeline
    with _dia_lock:
        if _dia_pipeline is None:
            print("[*] Loading pyannote diarization pipeline...")
            _dia_pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=hf_token,
            )
            if _DEVICE == "cuda" and torch.cuda.is_available():
                _dia_pipeline.to(torch.device("cuda"))
                print("[*] Diarization pipeline on GPU")
            else:
                print("[*] Diarization pipeline on CPU")
        return _dia_pipeline


def align_speakers(segments_data: list, diarization) -> dict:
    speaker_map = {}
    counter = [0]
    def label(spk):
        if spk not in speaker_map:
            speaker_map[spk] = chr(65 + counter[0])
            counter[0] += 1
        return speaker_map[spk]
    assignments = {}
    for idx, seg in enumerate(segments_data):
        start, end = seg['start'], seg['end']
        best_spk, best_overlap = None, 0.0
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            overlap = min(turn.end, end) - max(turn.start, start)
            if overlap > best_overlap:
                best_overlap = overlap
                best_spk = speaker
        assignments[idx] = label(best_spk) if best_spk else '?'
    return assignments


def run_transcription(task_id: str, filepath: str, model_size: str,
                      initial_prompt: str = '', hf_token: str = ''):
    q = _task_queues[task_id]
    segments_data = []
    try:
        q.put({'type': 'status', 'msg': f'Dang load model {model_size}...'})
        model = get_model(model_size)
        q.put({'type': 'status', 'msg': 'Dang phan tich audio...'})
        kwargs = dict(language='vi', beam_size=5, vad_filter=True,
                      vad_parameters=dict(min_silence_duration_ms=300, speech_pad_ms=200))
        if initial_prompt:
            kwargs['initial_prompt'] = initial_prompt
        segments, info = model.transcribe(filepath, **kwargs)
        q.put({'type': 'info', 'duration': round(info.duration, 1),
               'language': info.language, 'probability': round(info.language_probability, 2)})
        for idx, seg in enumerate(segments):
            sd = {'start': round(seg.start, 1), 'end': round(seg.end, 1), 'text': seg.text.strip()}
            segments_data.append(sd)
            q.put({'type': 'segment', 'idx': idx, **sd})
        q.put({'type': 'done', 'total': len(segments_data)})
        if hf_token and DIARIZATION_AVAILABLE and segments_data:
            q.put({'type': 'status', 'msg': 'Nhan dien giong noi (diarization)...'})
            try:
                pipeline = get_dia_pipeline(hf_token)
                diarization = pipeline(filepath)
                assignments = align_speakers(segments_data, diarization)
                num_speakers = len(set(assignments.values()) - {'?'})
                q.put({'type': 'speakers', 'assignments': assignments, 'count': num_speakers})
            except Exception as e:
                q.put({'type': 'dia_error', 'msg': str(e)})
    except Exception as e:
        q.put({'type': 'error', 'msg': str(e)})
    finally:
        try: os.remove(filepath)
        except: pass


@app.route('/')
def index():
    return render_template('index.html', diarization_available=DIARIZATION_AVAILABLE)

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'Empty filename'}), 400
    allowed = {'.mp3', '.mp4', '.m4a', '.wav', '.ogg', '.flac', '.webm', '.mkv'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        return jsonify({'error': f'Dinh dang khong ho tro: {ext}'}), 400
    task_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_FOLDER, f"{task_id}{ext}")
    file.save(filepath)
    return jsonify({'task_id': task_id, 'filename': file.filename,
                    'size': os.path.getsize(filepath), 'path': filepath})

@app.route('/transcribe', methods=['POST'])
def transcribe():
    data = request.json or {}
    task_id        = data.get('task_id')
    filepath       = data.get('path')
    model_size     = data.get('model', 'medium')
    initial_prompt = data.get('initial_prompt', '')
    hf_token       = data.get('hf_token', '').strip()
    if not task_id or not filepath or not os.path.exists(filepath):
        return jsonify({'error': 'Invalid task or file not found'}), 400
    q = queue.Queue()
    _task_queues[task_id] = q
    threading.Thread(target=run_transcription,
                     args=(task_id, filepath, model_size, initial_prompt, hf_token),
                     daemon=True).start()
    return jsonify({'task_id': task_id, 'started': True,
                    'diarization': bool(hf_token and DIARIZATION_AVAILABLE)})

@app.route('/stream/<task_id>')
def stream(task_id):
    def generate():
        q = _task_queues.get(task_id)
        if not q:
            yield f"data: {json.dumps({'type': 'error', 'msg': 'Task not found'})}\n\n"
            return
        while True:
            try:
                msg = q.get(timeout=60)
                yield f"data: {json.dumps(msg, ensure_ascii=False)}\n\n"
                if msg['type'] in ('error', 'speakers', 'dia_error'):
                    _task_queues.pop(task_id, None); break
            except queue.Empty:
                yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"
    return Response(generate(), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

@app.route('/export/docx', methods=['POST'])
def export_docx():
    data = request.json or {}
    segments = data.get('segments', [])
    filename = data.get('filename', 'transcript')
    has_speakers = any(seg.get('speaker') for seg in segments)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3); section.right_margin = Cm(2.5)
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TRANSCRIPT — PHIEN CHUYEN NGU')
    run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x10, 0x25, 0x90)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f'File: {filename}     Ngay: {datetime.now().strftime("%d/%m/%Y %H:%M")}     Doan: {len(segments)}')
    meta_run.font.size = Pt(9); meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()
    for seg in segments:
        start = _fmt_time(seg.get('start', 0)); end = _fmt_time(seg.get('end', 0))
        text = seg.get('text', ''); speaker = seg.get('speaker', '')
        p = doc.add_paragraph()
        if has_speakers and speaker:
            spk_run = p.add_run(f'[{speaker}]  ')
            spk_run.font.size = Pt(10); spk_run.font.bold = True
            spk_run.font.color.rgb = RGBColor(0x10, 0x25, 0x90)
        ts_run = p.add_run(f'[{start} -> {end}]  ')
        ts_run.font.size = Pt(9); ts_run.font.bold = True
        ts_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        txt_run = p.add_run(text); txt_run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    safe_name = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=safe_name)

def _fmt_time(sec: float) -> str:
    return f"{int(sec//60)}:{int(sec%60):02d}"

@app.route('/models')
def list_models():
    return jsonify({'loaded': list(_model_cache.keys()), 'diarization': DIARIZATION_AVAILABLE})

if __name__ == '__main__':
    print("=" * 50)
    print("  Vietnamese STT Web App")
    print("  http://localhost:5000")
    print("=" * 50)
    app.run(debug=False, host='0.0.0.0', port=5000, threaded=True)
